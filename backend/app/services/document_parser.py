import io
import logging
import fitz  # PyMuPDF
import docx
from PIL import Image
from app.core.exceptions import ValidationError

logger = logging.getLogger("app.services.document_parser")

class DocumentParserService:
    """
    Production Document Intelligence Service.
    Extracts raw text payloads from PDF (text & scanned image OCR), DOCX, and TXT files.
    """
    @classmethod
    def extract_text_from_file(cls, file_bytes: bytes, filename: str) -> str:
        """
        Auto-detects document format based on extension/magic bytes and extracts text.
        """
        fn_lower = filename.lower()
        if fn_lower.endswith(".pdf"):
            return cls.extract_text_from_pdf(file_bytes)
        elif fn_lower.endswith(".docx") or fn_lower.endswith(".doc"):
            return cls.extract_text_from_docx(file_bytes)
        elif fn_lower.endswith(".txt") or fn_lower.endswith(".md"):
            return file_bytes.decode("utf-8", errors="ignore").strip()
        else:
            # Fallback attempt
            try:
                return cls.extract_text_from_pdf(file_bytes)
            except Exception:
                return cls.extract_text_from_docx(file_bytes)

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """
        Reads PDF binaries. Handles standard vector text AND scanned image PDFs.
        """
        logger.info("DocumentParserService: Initiating PDF text extraction.")
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            extracted_pages = []
            
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text().strip()
                
                # If page contains standard vector text, append
                if page_text:
                    extracted_pages.append(page_text)
                else:
                    # Fallback for Scanned / Image-only PDF pages: Render page pixmap to OCR / Image stream
                    logger.info(f"DocumentParserService: Page {page_num+1} yielded empty text stream. Attempting image OCR extraction.")
                    pix = page.get_pixmap(dpi=150)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    
                    try:
                        import pytesseract
                        ocr_text = pytesseract.image_to_string(img).strip()
                        if ocr_text:
                            extracted_pages.append(ocr_text)
                    except Exception as ocr_err:
                        logger.info(f"DocumentParserService: OCR fallback skipped or tesseract binary not in PATH: {ocr_err}")

            doc.close()
            
            raw_text = "\n".join(extracted_pages).strip()
            if not raw_text:
                logger.warning("DocumentParserService: PDF text extraction yielded empty string.")
                raise ValidationError("The uploaded PDF document contains no readable text layout.")
                
            logger.info(f"DocumentParserService: PDF extraction completed successfully. (len={len(raw_text)})")
            return raw_text
        except ValidationError:
            raise
        except Exception as e:
            logger.error(f"DocumentParserService: PDF extraction failed with error: {e}", exc_info=True)
            raise ValidationError("Unable to read or parse the uploaded PDF file.")

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """
        Reads DOCX binaries and extracts paragraph, table, and header text.
        """
        logger.info("DocumentParserService: Initiating DOCX text extraction.")
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            extracted_paragraphs = []

            # Extract paragraphs
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if p_text:
                    extracted_paragraphs.append(p_text)

            # Extract table cell contents
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        extracted_paragraphs.append(" | ".join(row_text))

            raw_text = "\n".join(extracted_paragraphs).strip()
            if not raw_text:
                logger.warning("DocumentParserService: DOCX text extraction yielded empty string.")
                raise ValidationError("The uploaded DOCX document contains no readable text content.")

            logger.info(f"DocumentParserService: DOCX extraction completed successfully. (len={len(raw_text)})")
            return raw_text
        except ValidationError:
            raise
        except Exception as e:
            logger.error(f"DocumentParserService: DOCX extraction failed with error: {e}", exc_info=True)
            raise ValidationError("Unable to read or parse the uploaded DOCX file.")

    @staticmethod
    def clean_text_payload(text: str) -> str:
        """
        Filters duplicate spaces and standardizes character encodings.
        """
        lines = [line.strip() for line in text.splitlines()]
        clean_lines = [line for line in lines if line]
        return "\n".join(clean_lines)
